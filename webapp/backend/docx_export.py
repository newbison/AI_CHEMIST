"""Markdown → Word (.docx) 转换模块 — 专业排版版。

特性：
- 封面页（标题 + 副标题 + 日期）
- 颜色主题（深蓝主色 + 灰色辅助）
- 标题样式分级（H1 蓝底白字 / H2 深蓝 / H3 中蓝）
- 表格专业样式（表头深蓝底白字，斑马纹）
- 页眉页脚（报告标题 + 页码）
- 目录占位
"""
from __future__ import annotations

import io
import re
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


# 颜色主题
COLOR_PRIMARY = RGBColor(0x1A, 0x3C, 0x6E)      # 深蓝（主色）
COLOR_SECONDARY = RGBColor(0x2E, 0x5C, 0x8A)    # 中蓝（副色）
COLOR_ACCENT = RGBColor(0x0D, 0x94, 0x88)       # 青绿（强调）
COLOR_TEXT = RGBColor(0x1F, 0x29, 0x37)         # 正文
COLOR_MUTED = RGBColor(0x6B, 0x72, 0x80)        # 次要文字
COLOR_TABLE_HEADER_BG = "1A3C6E"                # 表头背景（hex 字符串）
COLOR_TABLE_ZEBRA = "F0F4F8"                    # 斑马纹背景

# 中文字体（Linux 部署用 Noto Sans CJK SC，Windows 用 Microsoft YaHei）
# 可通过 CN_FONT 环境变量覆盖
import os
CN_FONT = os.environ.get("CN_FONT", "Noto Sans CJK SC" if os.name != "nt" else "Microsoft YaHei")


def _set_run_font(run, *, size: float | None = None, bold: bool = False,
                  color: RGBColor | None = None, font_name: str = CN_FONT) -> None:
    """统一设置 run 字体，同时处理中文 eastAsia 字体。

    python-docx 只设 font.name 对中文不生效，必须额外设 w:eastAsia。
    """
    run.font.name = font_name
    # 设置中文字体（eastAsia）
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn('w:rFonts'))
    if r_fonts is None:
        r_fonts = OxmlElement('w:rFonts')
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn('w:eastAsia'), CN_FONT)
    r_fonts.set(qn('w:ascii'), font_name)
    r_fonts.set(qn('w:hAnsi'), font_name)
    if size is not None:
        run.font.size = Pt(size)
    if bold:
        run.font.bold = True
    if color is not None:
        run.font.color.rgb = color


def _set_cell_bg(cell, color_hex: str) -> None:
    """设置单元格背景色。"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def _set_cell_border(cell, color_hex: str = "CCCCCC", size: str = "4") -> None:
    """设置单元格边框。"""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), size)
        border.set(qn('w:color'), color_hex)
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def _add_page_break(doc) -> None:
    """添加分页符。"""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def _add_cover_page(doc, title: str, subtitle: str = "") -> None:
    """添加封面页。"""
    # 顶部留白
    for _ in range(6):
        doc.add_paragraph()

    # 主标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    _set_run_font(run, size=28, bold=True, color=COLOR_PRIMARY)

    # 副标题
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitle)
        _set_run_font(run, size=14, color=COLOR_SECONDARY)

    # 分隔线
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('━' * 20)
    _set_run_font(run, size=12, color=COLOR_ACCENT)

    # 日期
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'生成日期：{datetime.now().strftime("%Y-%m-%d")}')
    _set_run_font(run, size=11, color=COLOR_MUTED)

    _add_page_break(doc)


def _style_heading(paragraph, level: int) -> None:
    """为标题段落应用样式。"""
    for run in paragraph.runs:
        if level == 1:
            _set_run_font(run, size=18, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        elif level == 2:
            _set_run_font(run, size=15, bold=True, color=COLOR_PRIMARY)
        elif level == 3:
            _set_run_font(run, size=13, bold=True, color=COLOR_SECONDARY)
        else:
            _set_run_font(run, size=12, bold=True, color=COLOR_TEXT)


def _add_h1_with_bg(doc, text: str) -> None:
    """H1 标题：深蓝底白字。"""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    _set_cell_bg(cell, COLOR_TABLE_HEADER_BG)
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    _set_run_font(run, size=18, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    # 去掉单元格边框
    _set_cell_border(cell, "1A3C6E", "0")
    doc.add_paragraph()


def _add_inline_formatting(paragraph, text: str) -> None:
    """处理行内格式：**加粗**、`代码`、[链接](url)。"""
    pattern = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))')
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            _set_run_font(run, bold=True, color=COLOR_PRIMARY)
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            _set_run_font(run, size=10, font_name='Consolas', color=COLOR_ACCENT)
        elif part.startswith('[') and '](' in part:
            m = re.match(r'\[([^\]]+)\]\(([^)]+)\)', part)
            if m:
                run = paragraph.add_run(m.group(1))
                _set_run_font(run, color=COLOR_SECONDARY)
                run.underline = True
            else:
                run = paragraph.add_run(part)
                _set_run_font(run, color=COLOR_TEXT)
        else:
            run = paragraph.add_run(part)
            _set_run_font(run, color=COLOR_TEXT)


def _parse_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith('|'):
        line = line[1:]
    if line.endswith('|'):
        line = line[:-1]
    return [cell.strip() for cell in line.split('|')]


def _is_table_separator(line: str) -> bool:
    return bool(re.match(r'^\s*\|?[\s\-:]+(\|[\s\-:]+)+\|?\s*$', line))


def _add_styled_table(doc, headers: list[str], rows: list[list[str]]) -> None:
    """添加专业样式表格：深蓝表头 + 斑马纹。"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        _set_cell_bg(cell, COLOR_TABLE_HEADER_BG)
        _set_cell_border(cell, "1A3C6E", "4")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        _set_run_font(run, size=9.5, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    # 数据行
    for ri, row in enumerate(rows):
        is_zebra = ri % 2 == 1
        for j, cell_text in enumerate(row):
            if j >= len(headers):
                continue
            cell = table.rows[ri + 1].cells[j]
            if is_zebra:
                _set_cell_bg(cell, COLOR_TABLE_ZEBRA)
            _set_cell_border(cell, "D0D7DE", "4")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ''
            p = cell.paragraphs[0]
            _add_inline_formatting(p, cell_text)
            for run in p.runs:
                _set_run_font(run, size=9, color=run.font.color.rgb or COLOR_TEXT)

    doc.add_paragraph()


def _add_footer(section, report_title: str) -> None:
    """添加页脚：报告标题 + 页码。"""
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 报告标题
    run = p.add_run(f'{report_title}  |  ')
    _set_run_font(run, size=8, color=COLOR_MUTED)

    # 页码字段
    fld_char1 = OxmlElement('w:fldChar')
    fld_char1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'PAGE'
    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'end')

    run = p.add_run()
    _set_run_font(run, size=8, color=COLOR_MUTED)
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def markdown_to_docx(markdown: str, *, title: str | None = None) -> bytes:
    """把 Markdown 文本转成 .docx 字节流。

    Args:
        markdown: Markdown 格式报告文本
        title: 封面标题，默认从 Markdown 第一个 H1 提取
    """
    doc = Document()

    # 设置默认字体（含中文 eastAsia）
    style = doc.styles['Normal']
    style.font.name = CN_FONT
    style.font.size = Pt(10.5)
    style.font.color.rgb = COLOR_TEXT
    # Normal 样式也要设 eastAsia，否则正文中文不生效
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.find(qn('w:rFonts'))
    if r_fonts is None:
        r_fonts = OxmlElement('w:rFonts')
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn('w:eastAsia'), CN_FONT)
    r_fonts.set(qn('w:ascii'), CN_FONT)
    r_fonts.set(qn('w:hAnsi'), CN_FONT)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(4)

    # 页边距
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 提取标题
    if not title:
        first_heading = re.search(r'^#\s+(.+)$', markdown, re.MULTILINE)
        title = first_heading.group(1).strip() if first_heading else "R&D 智能报告"

    # 封面
    _add_cover_page(doc, title, "R&D Portfolio Intelligence Report")

    # 页脚
    _add_footer(doc.sections[0], title)

    # 解析正文
    lines = markdown.split('\n')
    i = 0
    in_code_block = False
    code_buffer: list[str] = []
    skip_first_h1 = True  # 跳过第一个 H1（已用作封面标题）

    while i < len(lines):
        line = lines[i]

        # 代码块
        if line.strip().startswith('```'):
            if in_code_block:
                p = doc.add_paragraph()
                run = p.add_run('\n'.join(code_buffer))
                _set_run_font(run, size=9, font_name='Consolas', color=COLOR_ACCENT)
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # 空行
        if not line.strip():
            i += 1
            continue

        # 标题
        heading_match = re.match(r'^(#{1,4})\s+(.*)', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            # 跳过第一个 H1
            if level == 1 and skip_first_h1:
                skip_first_h1 = False
                i += 1
                continue
            if level == 1:
                _add_h1_with_bg(doc, text)
            else:
                heading = doc.add_heading(text, level=min(level, 4))
                _style_heading(heading, level)
            i += 1
            continue

        # 分隔线
        if line.strip() in ('---', '***', '___'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run('─' * 40)
            _set_run_font(run, size=9, color=COLOR_MUTED)
            i += 1
            continue

        # 引用
        if line.strip().startswith('>'):
            text = re.sub(r'^>\s*', '', line.strip())
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run(text)
            _set_run_font(run, size=10, color=COLOR_MUTED)
            run.italic = True
            i += 1
            continue

        # 表格
        if '|' in line and i + 1 < len(lines) and _is_table_separator(lines[i + 1]):
            table_lines = [line]
            i += 1  # 跳过分隔行
            i += 1
            while i < len(lines) and '|' in lines[i] and lines[i].strip():
                table_lines.append(lines[i])
                i += 1
            headers = _parse_table_row(table_lines[0])
            rows = [_parse_table_row(r) for r in table_lines[1:]]
            _add_styled_table(doc, headers, rows)
            continue

        # 无序列表
        list_match = re.match(r'^(\s*)[-*+]\s+(.*)', line)
        if list_match:
            indent = len(list_match.group(1))
            text = list_match.group(2)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25 + indent * 0.2)
            _add_inline_formatting(p, text)
            i += 1
            continue

        # 有序列表
        ol_match = re.match(r'^(\s*)\d+\.\s+(.*)', line)
        if ol_match:
            indent = len(ol_match.group(1))
            text = ol_match.group(2)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.left_indent = Inches(0.25 + indent * 0.2)
            _add_inline_formatting(p, text)
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        _add_inline_formatting(p, line.strip())
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
